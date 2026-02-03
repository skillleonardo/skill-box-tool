import docx
from docx.shared import Pt
import os

source_path = "CODCORSO_audio.docx"
target_path = "affidotransformer/assets/template.docx"

if not os.path.exists(source_path):
    print("Source file not found!")
    exit(1)

doc = docx.Document(source_path)

# 1. Replace placeholders in Paragraphs
def replace_in_paragraphs(doc):
    replacements = {
        "codice corso01": "{{codice_corso}}",
        "Nome cliente01": "{{nome_cliente}}",
        "Data Corrente01": "{{data_corrente}}"
    }
    for p in doc.paragraphs:
        for key, val in replacements.items():
            if key in p.text:
                # Naive replacement in runs
                # Ideally we clear runs and set text, but that removes formatting.
                # Since the keys are unique enough, we can try string replacement
                # converting the whole paragraph text might lose bold/italic inside the key?
                # The keys seem simple.
                # A safer way in python-docx:
                p.text = p.text.replace(key, val)

replace_in_paragraphs(doc)

# 2. Handle Tables
# We want to keep Table 0 (Metadata).
# We want to use Table 1 as the main list.
# We remove Table 2, 3, 4 ... if they exist.

tables = doc.tables
# Ensure we have at least 2 tables
if len(tables) < 2:
    print("Not enough tables!")
    exit(1)

# Remove tables backwards from end until index 1
for i in range(len(tables) - 1, 1, -1):
    t = tables[i]
    t._element.getparent().remove(t._element)

# Now edit Table 1
table = doc.tables[1]
# Keep Header (Row 0)
# Delete all other rows
for i in range(len(table.rows) - 1, 0, -1):
    row = table.rows[i]
    row._element.getparent().remove(row._element)

# Add the Loop Row
row = table.add_row()
# We need to ensure the cells match styles if possible
# Set up the loop
# Column 0: {#items}{filename}
# Column 1: {content}{/items}
# The text needs to be in runs to work well, but docxtemplater is string based usually.

row.cells[0].text = "{#items}{filename}"
row.cells[1].text = "{content}{/items}"

# Optional: Set font/style if needed? 
# We assume the user wants default table style.

# Save
doc.save(target_path)
print(f"Template saved to {target_path}")
