import docx
import os

doc_path = "CODCORSO_audio.docx"
doc = docx.Document(doc_path)

with open("analysis_result.txt", "w", encoding="utf-8") as f:
    f.write("Searching tables for 'Nome file'...\n")
    for i, table in enumerate(doc.tables):
        f.write(f"--- Table {i} ---\n")
        for r_idx, row in enumerate(table.rows):
            row_text = [c.text.strip() for c in row.cells]
            f.write(f"Row {r_idx}: {row_text}\n")
            if r_idx > 10: # Limit rows
                f.write("...\n")
                break
    
    f.write("\n--- Paragraphs ---\n")
    for p in doc.paragraphs:
        if "codice corso" in p.text.lower() or "nome cliente" in p.text.lower() or "data corrente" in p.text.lower():
             f.write(f"PARA: {p.text}\n")
